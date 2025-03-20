import sys
import os
from fastapi import FastAPI, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import tempfile
from datetime import datetime
import json
import httpx

# Add the current directory to the Python path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Try to import OpenAI for AI-based generation
try:
    from dotenv import load_dotenv
    load_dotenv()
    
    # Print debugging information
    print("Environment variables loaded from .env file")
    openai_api_key = os.getenv("OPENAI_API_KEY")
    print(f"OpenAI API Key available: {bool(openai_api_key)}")
    
    if openai_api_key:
        from openai import OpenAI
        client = OpenAI(
            api_key=openai_api_key,
            http_client=httpx.Client()
        )
        has_openai = True
        print("OpenAI client initialized successfully")
    else:
        has_openai = False
        print("ERROR: OpenAI API key not found in environment variables")
except ImportError as e:
    has_openai = False
    print(f"ERROR: Could not import required packages: {e}")
    print("Please install them with: pip install python-dotenv openai")

# Variable to store the last generated priorities and organization name
last_generated = {
    "org_name": "",
    "org_website": "",
    "priorities": []
}

class OrgData(BaseModel):
    org_name: str
    org_website: str

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174",
                   "https://strategic-priorities-frontend.onrender.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "Welcome to the Strategic Priorities Generator API"}

def generate_ai_priorities(org_name, org_website):
    """Generate strategic priorities using OpenAI."""
    if not has_openai:
        print("OpenAI functionality not available")
        return None
    
    try:
        print(f"Generating priorities for {org_name} using OpenAI...")
        
        prompt = f"""
        Please identify EXACTLY 5 strategic priorities for {org_name} based on its website: {org_website}.
        You MUST provide EXACTLY 5 priorities, each with at least 5 initiatives.

        Each priority should include:
        - A clear title
        - A description explaining why this priority exists
        - Exactly 5 result definitions that describe how this priority is achieved.

        Please consider:
        1. Use information from {org_website} if possible, but also use your general knowledge about city governance and best practices for municipalities similar to {org_name}.
        2. Include both community-oriented priorities (like public safety, economic development) as well as internal governance priorities (like fiscal responsibility, transparent government).
        3. Be descriptive and detailed in your explanations.
        4. When possible, include the specific source or page from {org_website} where you found the information.

        The output should follow this JSON structure precisely:
        [
          {{
            "priority": "Priority Title",
            "description": "Description of the priority",
            "definitions": [
              {{
                "title": "Result Definition Title",
                "description": "Description of how this result is achieved",
                "source": "URL or section of the website where this information was found (if available)"
              }},
              ... (4 more definitions for each priority)
            ]
          }},
          ... (4 more priorities for a total of EXACTLY 5)
        ]

        Return ONLY the JSON structure with no additional explanation.
        """
        
        # Call OpenAI API with system message and increased max_tokens
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a strategic planning expert who always provides exactly 5 strategic priorities with exactly 5 initiatives each when asked."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3500  # Increased to allow for more content
        )
        
        # Extract content from response
        content = response.choices[0].message.content.strip()
        print(f"OpenAI response received: {len(content)} characters")
        
        # Parse the JSON content
        try:
            # Find the start and end of the JSON structure
            json_start = content.find('[')
            json_end = content.rfind(']') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_content = content[json_start:json_end]
                priorities = json.loads(json_content)
                print(f"Successfully parsed JSON with {len(priorities)} priorities")
                return priorities
            else:
                # If we can't find valid JSON brackets, try parsing the whole response
                priorities = json.loads(content)
                print(f"Successfully parsed full response as JSON with {len(priorities)} priorities")
                return priorities
                
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON from OpenAI response: {e}")
            print(f"Response content: {content}")
            return None
                
    except Exception as e:
        print(f"Error generating priorities with OpenAI: {e}")
        return None

@app.post("/generate")
async def generate_priorities_endpoint(data: OrgData):
    """Generate strategic priorities for an organization."""
    global last_generated
    
    # Store the organization info for later use in downloads
    last_generated["org_name"] = data.org_name
    last_generated["org_website"] = data.org_website
    
    # Try AI generation first (if available)
    ai_priorities = None
    if has_openai:
        ai_priorities = generate_ai_priorities(data.org_name, data.org_website)
    
    # If AI generation fails or isn't available, use mock data
    if ai_priorities is None:
        print("Using mock priorities data")
        priorities = [
            {
                "priority": "Safe and Secure Communities",
                "description": f"Ensuring public safety and justice for all residents of {data.org_name} by protecting communities from crime and harm while upholding fairness and accountability in the justice system.",
                "definitions": [
                    {
                        "title": "Effective & Equitable Law Enforcement:",
                        "description": "Prevent crime and respond quickly to emergencies through community-oriented policing and accountable practices that build public trust.",
                        "source": f"{data.org_website}/police-department"
                    },
                    {
                        "title": "Justice Reform & Alternatives to Incarceration:",
                        "description": "Implement diversion programs and treatment services (embodying the \"Care First, Jails Last\" approach) so social and mental health issues are addressed with health interventions instead of jail whenever possible.",
                        "source": f"{data.org_website}/courts"
                    },
                    {
                        "title": "Fire and Emergency Medical Services:",
                        "description": "Provide effective fire protection and emergency medical response across all communities, ensuring well-coordinated responses to fires, accidents, and medical emergencies to safeguard lives and property.",
                        "source": f"{data.org_website}/fire-department"
                    },
                    {
                        "title": "Emergency Preparedness & Disaster Response:",
                        "description": "Maintain robust preparedness plans and rapid-response protocols for natural disasters, public health crises, and other emergencies, coordinating law enforcement, fire, and health agencies to protect the public.",
                        "source": f"{data.org_website}/emergency-management"
                    },
                    {
                        "title": "Protecting Vulnerable Populations:",
                        "description": "Collaborate across agencies to prevent abuse, neglect, and exploitation, focusing on at-risk children, seniors, and other vulnerable groups through early intervention and supportive services.",
                        "source": f"{data.org_website}/community-services"
                    }
                ]
            },
            {
                "priority": "Economic Development & Job Creation",
                "description": f"Strengthen {data.org_name}'s economy by supporting business growth, attracting investment, and creating quality jobs that provide pathways to prosperity for all residents.",
                "definitions": [
                    {
                        "title": "Business Growth & Support:",
                        "description": "Facilitate the creation and expansion of local businesses through streamlined permitting, technical assistance, and access to capital, with particular focus on small and minority-owned enterprises.",
                        "source": f"{data.org_website}/economic-development"
                    },
                    {
                        "title": "Workforce Development:",
                        "description": "Partner with educational institutions and industries to provide job training, skill development, and career pathways that prepare residents for quality jobs in growing sectors.",
                        "source": f"{data.org_website}/workforce"
                    },
                    {
                        "title": "Strategic Investment:",
                        "description": "Target public investments to catalyze private development in key geographic areas and industry sectors, creating jobs and strengthening the tax base.",
                        "source": f"{data.org_website}/development-projects"
                    },
                    {
                        "title": "Tourism Promotion:",
                        "description": "Develop and market local attractions, events, and amenities to increase tourism and visitor spending in the local economy.",
                        "source": f"{data.org_website}/tourism"
                    },
                    {
                        "title": "Regional Collaboration:",
                        "description": "Work with neighboring communities and regional economic development organizations to attract industry and create a stronger overall economic ecosystem.",
                        "source": f"{data.org_website}/regional-partnerships"
                    }
                ]
            },
            {
                "priority": "Infrastructure & Sustainable Development",
                "description": f"Develop and maintain {data.org_name}'s physical infrastructure to support quality of life, economic vitality, and environmental sustainability.",
                "definitions": [
                    {
                        "title": "Transportation Networks:",
                        "description": "Develop and maintain a comprehensive, multimodal transportation system that safely and efficiently moves people and goods while reducing congestion and environmental impacts.",
                        "source": f"{data.org_website}/transportation"
                    },
                    {
                        "title": "Sustainable Environmental Practices:",
                        "description": "Implement policies and programs that conserve natural resources, reduce pollution, and build resilience to climate change impacts.",
                        "source": f"{data.org_website}/sustainability"
                    },
                    {
                        "title": "Smart Growth & Planning:",
                        "description": "Guide development to create livable, walkable communities that balance housing, jobs, and services while preserving open space and community character.",
                        "source": f"{data.org_website}/planning"
                    },
                    {
                        "title": "Utility Services & Infrastructure:",
                        "description": "Ensure reliable, efficient utility services including water, sewer, and waste management that meet current needs and accommodate future growth.",
                        "source": f"{data.org_website}/utilities"
                    },
                    {
                        "title": "Parks & Public Spaces:",
                        "description": "Develop and maintain parks, trails, and public spaces that enhance quality of life, promote active living, and protect natural resources.",
                        "source": f"{data.org_website}/parks"
                    }
                ]
            },
            {
                "priority": "Responsive & Effective Governance",
                "description": f"Deliver high-quality public services through transparent, accountable, and fiscally responsible government operations that engage the community and respond to residents' needs.",
                "definitions": [
                    {
                        "title": "Fiscal Responsibility:",
                        "description": "Manage public resources efficiently and effectively through sound budgeting, responsible financial planning, and transparent reporting.",
                        "source": f"{data.org_website}/finance"
                    },
                    {
                        "title": "Community Engagement:",
                        "description": "Involve residents in government decision-making through meaningful public participation opportunities, accessible information, and responsive communication.",
                        "source": f"{data.org_website}/community-engagement"
                    },
                    {
                        "title": "Technology & Innovation:",
                        "description": "Leverage technology to improve service delivery, increase efficiency, and enhance communication with residents.",
                        "source": f"{data.org_website}/technology"
                    },
                    {
                        "title": "Workforce Excellence:",
                        "description": "Recruit, develop, and retain a skilled, diverse public workforce committed to high-quality service delivery and continuous improvement.",
                        "source": f"{data.org_website}/jobs"
                    },
                    {
                        "title": "Intergovernmental Collaboration:",
                        "description": "Work effectively with other levels of government and neighboring jurisdictions to address shared challenges and maximize resources.",
                        "source": f"{data.org_website}/government"
                    }
                ]
            },
            {
                "priority": "Quality of Life & Community Wellbeing",
                "description": f"Enhance the overall quality of life in {data.org_name} by promoting community health, expanding cultural and recreational opportunities, and ensuring access to quality housing and essential services.",
                "definitions": [
                    {
                        "title": "Housing Affordability & Access:",
                        "description": "Support the development and preservation of diverse, quality housing options accessible to households of all income levels.",
                        "source": f"{data.org_website}/housing"
                    },
                    {
                        "title": "Arts, Culture & Recreation:",
                        "description": "Provide and support a wide range of arts, cultural, and recreational programs and facilities that enhance community identity and quality of life.",
                        "source": f"{data.org_website}/recreation"
                    },
                    {
                        "title": "Public Health & Wellness:",
                        "description": "Promote physical and mental health through preventive health services, healthy environment initiatives, and expanded access to healthcare.",
                        "source": f"{data.org_website}/health"
                    },
                    {
                        "title": "Education & Lifelong Learning:",
                        "description": "Support quality education and lifelong learning opportunities through partnerships with schools, libraries, and other educational organizations.",
                        "source": f"{data.org_website}/education"
                    },
                    {
                        "title": "Diversity & Inclusion:",
                        "description": "Foster a welcoming, inclusive community that celebrates diversity, promotes equity, and ensures all residents have opportunities to thrive.",
                        "source": f"{data.org_website}/diversity"
                    }
                ]
            }
        ]
    else:
        # Use the AI-generated priorities
        priorities = ai_priorities
    
    # Store the priorities for later use in downloads
    last_generated["priorities"] = priorities
    
    return {"priorities": priorities}

def create_simple_word_doc(priorities, org_name):
    """Create a simple Word document with the strategic priorities."""
    try:
        from docx import Document
        doc = Document()
        
        # Add a title
        doc.add_heading(f'Strategic Priorities for {org_name}', 0)
        
        # Add a timestamp
        doc.add_paragraph(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add each priority
        for i, priority in enumerate(priorities):
            doc.add_heading(f'Priority {i+1}: {priority["priority"]}', 1)
            doc.add_paragraph(priority["description"])
            
            doc.add_heading('Key Initiatives:', 2)
            for definition in priority["definitions"]:
                p = doc.add_paragraph()
                p.add_run(f'{definition["title"]}').bold = True
                p.add_run(f' {definition["description"]}')
                
                # Add source if available
                if "source" in definition and definition["source"]:
                    source_p = doc.add_paragraph(f'Source: {definition["source"]}')
                    source_p.style = 'Caption'
        
        # Save the document to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
    except ImportError:
        print("python-docx library not found. Please install it with 'pip install python-docx'")
        return None

def create_simple_excel(priorities, org_name):
    """Create a simple Excel file with the strategic priorities."""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Strategic Priorities"
        
        # Add a title
        ws['A1'] = f'Strategic Priorities for {org_name}'
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:E1')
        
        # Add a timestamp
        ws['A2'] = f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws.merge_cells('A2:E2')
        
        # Add headers
        headers = ['Priority', 'Description', 'Initiative', 'Details', 'Source']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.font = Font(bold=True)
        
        # Add data
        row = 4
        for priority in priorities:
            for i, definition in enumerate(priority['definitions']):
                if i == 0:
                    # First row of a priority includes priority title and description
                    ws.cell(row=row, column=1, value=priority['priority'])
                    ws.cell(row=row, column=2, value=priority['description'])
                else:
                    # Subsequent rows of the same priority have empty cells for priority and description
                    ws.cell(row=row, column=1, value='')
                    ws.cell(row=row, column=2, value='')
                
                # Add initiative details
                ws.cell(row=row, column=3, value=definition['title'])
                ws.cell(row=row, column=4, value=definition['description'])
                
                # Add source if available
                if "source" in definition:
                    ws.cell(row=row, column=5, value=definition.get('source', ''))
                
                row += 1
        
        # Auto-adjust column widths
        for col in range(1, 6):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 30
        
        # Save to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        wb.save(temp_file.name)
        return temp_file.name
    except ImportError:
        print("openpyxl library not found. Please install it with 'pip install openpyxl'")
        return None

@app.get("/download/word")
async def download_word_endpoint():
    """Download strategic priorities as a Word document."""
    global last_generated
    
    try:
        # Use the stored organization name and priorities
        org_name = last_generated["org_name"]
        if not org_name:
            org_name = "Your Organization"
            
        priorities = last_generated["priorities"]
        if not priorities:
            # Fallback data if nothing has been generated
            priorities = [
                {
                    "priority": "Example Priority",
                    "description": "This is an example priority. Please generate priorities first.",
                    "definitions": [
                        {
                            "title": "Example Initiative:",
                            "description": "This is an example initiative.",
                            "source": "Example source"
                        }
                    ]
                }
            ]
        
        print(f"Creating Word document for '{org_name}' with {len(priorities)} priorities")
        file_path = create_simple_word_doc(priorities, org_name)
        
        if file_path:
            return FileResponse(
                path=file_path,
                filename="strategic_priorities.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            return {"error": "Failed to create Word document. Make sure python-docx is installed."}
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return {"error": "Failed to create Word document"}

@app.get("/download/excel")
async def download_excel_endpoint():
    """Download strategic priorities as an Excel spreadsheet."""
    global last_generated
    
    try:
        # Use the stored organization name and priorities
        org_name = last_generated["org_name"]
        if not org_name:
            org_name = "Your Organization"
            
        priorities = last_generated["priorities"]
        if not priorities:
            # Fallback data if nothing has been generated
            priorities = [
                {
                    "priority": "Example Priority",
                    "description": "This is an example priority. Please generate priorities first.",
                    "definitions": [
                        {
                            "title": "Example Initiative:",
                            "description": "This is an example initiative.",
                            "source": "Example source"
                        }
                    ]
                }
            ]
        
        print(f"Creating Excel document for '{org_name}' with {len(priorities)} priorities")
        file_path = create_simple_excel(priorities, org_name)
        
        if file_path:
            return FileResponse(
                path=file_path,
                filename="strategic_priorities.xlsx",
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            return {"error": "Failed to create Excel file. Make sure openpyxl is installed."}
    except Exception as e:
        print(f"Error creating Excel file: {e}")
        return {"error": "Failed to create Excel file"}

if __name__ == "__main__":
    import uvicorn
    
    # Install required packages if not present
    missing_packages = []
    try:
        import docx
    except ImportError:
        missing_packages.append("python-docx")
    
    try:
        import openpyxl
    except ImportError:
        missing_packages.append("openpyxl")
    
    if missing_packages:
        print(f"Warning: Missing packages: {', '.join(missing_packages)}")
        print(f"Please install them with: pip install {' '.join(missing_packages)}")
    
    uvicorn.run(app, host="127.0.0.1", port=8000)