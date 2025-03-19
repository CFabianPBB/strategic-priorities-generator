from docx import Document

def create_word(priorities):
    doc = Document()
    doc.add_heading('Strategic Priorities', level=1)

    for priority in priorities:
        doc.add_heading(priority['priority'], level=2)
        doc.add_paragraph(priority['description'])
        doc.add_paragraph("Result Definitions:")
        for definition in priority['definitions']:
            doc.add_paragraph(f"- {definition}")

    file_path = "strategic_priorities.docx"
    doc.save(file_path)
    return file_path
